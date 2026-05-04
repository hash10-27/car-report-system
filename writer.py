from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from parser import fix_dtc
from docx.shared import RGBColor, Pt
import re

# 🔹 استبدال النص داخل الفقرات والجداول

def replace_all(doc, key, value):
    # الفقرات
    for p in doc.paragraphs:
        if key in p.text:
            p.text = p.text.replace(key, value)

    # الجداول  
    for table in doc.tables:  
        for row in table.rows:  
            for cell in row.cells:  
                if key in cell.text:  
                    cell.text = cell.text.replace(key, value)

def build_dtc_text(dtc_list):
    if not dtc_list:
        return "لا يوجد أعطال"

    lines = []  
    for d in dtc_list:  
        line = f"{d['system']} | {d['code']} | {d['desc']}"  
        lines.append(line)  

    return "\n".join(lines)

def center_cell(cell):
    # 🔥 الفقرة (المهم)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pPr = paragraph._element.get_or_add_pPr()

        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    # 🔽 الخلية (اختياري لكن جيد)
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
# 🔹 تحويل قائمة DTC إلى نص مرتب

def fill_dtc_table(doc, dtc_list):
    table = doc.tables[1]  # أول جدول في القالب  

    for d in dtc_list:  
        row_cells = table.add_row().cells  
        row_cells[0].text = d["system"]  
        row_cells[1].text = d["code"]  
        row_cells[2].text = d["desc"]

def format_systems(systems):
    if not systems:
        return "لا يوجد"
    return "\n".join(systems)


def fill_ok_systems_table(doc, systems_ok):
    # 📌 اختر الجدول المناسب (آخر جدول)  
    table = doc.tables[-1]  

    for system in systems_ok:  
        # ❌ فلتر نهائي  
        if any(x in system for x in [  
            "إخلاء المسؤولية",  
            "هذا التقرير",  
            "بيانات",  
            "LAUNCH"  
        ]):  
            continue  

        row = table.add_row().cells  
        clean_name = re.sub(r'^d+.', '', system).strip()  
        row[0].text = clean_name  
        row[1].text = "✔"  

        # تنسيق  
        style_cell(row[0])  
        style_cell(row[1], bold=True, color=RGBColor(0, 150, 0))  

        # 🔥 توسيط  
        for cell in row:  
            for p in cell.paragraphs:  
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
def style_cell(cell, bold=False, color=None):
    font_name = "Arial"  # 👈 استخدم Arial أو Segoe UI

    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(11)

            # 🔥 هذا المهم فعلاً
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')

            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:cs'), font_name)

            rPr.append(rFonts)

            if color:
                run.font.color.rgb = color

def extract_raw_dtc_block(text):
    lines = text.split("")
    capture = False
    result = []

    for line in lines:  
        if "رمز خطأ النظام" in line:  
            capture = True  
            continue  
        if "على ما يرام" in line:  
            break  
        if capture:  
            result.append(line.strip())  

    return "\n".join(result)

def build_dtc_text(dtc_list):
    if not dtc_list:
        return "لا يوجد أعطال"
    lines = []
    for d in dtc_list:
        system = d.get('system') or d.get('title') or 'غير محدد'
        line = f"{system} | {d.get('code','')} | {d.get('desc','')}"
        lines.append(line)
    return "\n".join(lines)

def clean_title(text):
    text = text.strip()

    # 🔥 فصل العربي عن الإنجليزي
    text = re.sub(r'([A-Za-z])([؀-ۿ])', r'\1 \2', text)
    text = re.sub(r'([؀-ۿ])([A-Za-z])', r'\1 \2', text)

    # 🔥 فصل الكلمات الكبيرة الملتصقة
    text = re.sub(r'([A-Z]{2,})([A-Z]{2,})', r'\1 \2', text)

    # 🔥 إزالة التكرار الغبي مثل EOBDOBD
    words = text.split()
    cleaned = []
    for w in words:
        if w not in cleaned:
            cleaned.append(w)

    text = " ".join(cleaned)

    # 🔥 ترتيب OBD إذا انعكس
    if "OBD" in text[::-1]:
        text = text[::-1]

    return text.strip()

def fill_system_tables(doc, faults_raw):
    seen_codes = set()
    table = doc.tables[1]
    current_title = ""
    current_group = False

    if isinstance(faults_raw, str):
        faults_raw = faults_raw.splitlines()

    for line in faults_raw:
        print(f"\n🔹 LINE: {line}") 

        line = line.strip()
        if not line:
            continue

        def has_dtc(line):
            return re.search(r'\d+\.\d+[A-Z0-9]{4}[PCBU]', line) or re.search(r'\d+\.[0-9A-Z]{4}[PCBU]', line)

        # 🔥 الحالة 1: عنوان
        if not has_dtc(line) or line.startswith("نظام"):

            if any(x in line for x in ["غير طبيعي", "DTC", "Present", "الحالي", "التاريخ"]):
                continue

            if current_title:
                if (
                    len(line) < 50
                    and not line.startswith("نظام")
                    and not re.match(r'DTC\s*\d+', line)
                    and not re.search(r'\([A-Z/ ]+\)', line)  # 🔥 هذا هو الحل
                ):
                    current_title += " " + line
                    continue

            current_title = clean_title(line)

            row = table.add_row().cells
            row[0].text = f"🔹 {current_title}"
            row[1].text = ""
            row[2].text = ""

            style_cell(row[0], bold=True, color=RGBColor(0, 70, 160))
            center_cell(row[0])
            continue

        # 🔥 الحالة 2: DTC
        parts = re.split(r'(?=\d+\.\d+[A-Z0-9]{4}[PCBU]|\d+\.[0-9A-Z]{4}[PCBU])', line)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            m = re.search(r'(\d+\.\d+[A-Z0-9]{4}[PCBU]|\d+\.[0-9A-Z]{4}[PCBU])', part)
            if not m:
                continue

            row = table.add_row().cells
            title_to_use = current_title or 'غير محدد'

            row[0].text = title_to_use if part == parts[0] else ""
            code_raw = m.group(0).split('.')[-1]
            code_fixed = fix_dtc(code_raw)

            # 🔥 منع التكرار
            key = code_fixed

            if key in seen_codes:
                continue

            seen_codes.add(key)

            row[1].text = code_fixed
            for p in row[1].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(200, 0, 0)  # 🔴 

            desc = part[m.end():].strip()

            # 🔥 إضافة هذا فقط
            next_index = faults_raw.index(line) + 1
            if next_index < len(faults_raw):
                next_line = faults_raw[next_index].strip()
                if next_line and not has_dtc(next_line):
                    desc += " " + next_line

            desc = re.sub(r'^(الحالي|التاريخ)\s*', '', desc)

            row[2].text = desc
            style_cell(row[2])

            style_cell(row[0], bold=True)
            center_cell(row[0])
            center_cell(row[1])
            center_cell(row[2])
# 🔹 تعبئة القالب

def fill_template(template_path, output_path, data):
    doc = Document(template_path)
    fill_ok_systems_table(doc, data["systems_ok"])

    replacements = {  
        "{year}": data["car_info"]["year"],  
        "{make}": data["car_info"]["make"],  
        "{model}": data["car_info"]["model"],  
        "{vin}": data["car_info"]["vin"],  
        "{engine}": data["car_info"]["engine"],  
        "{mileage}": data["car_info"]["mileage"],  

        "{customer}": data["customer_info"]["customer"],  
        "{technician}": data["customer_info"]["technician"],  
        "{phone}": data["customer_info"]["phone"],  

        "{car_version}": data["meta"]["car_version"],  
        "{app_version}": data["meta"]["app_version"],  
        "{test_time}": data["meta"]["test_time"],  
        "{sn}": data["meta"]["sn"],  
        "{systems_ok}": "\n".join(data["systems_ok"]),
        "{systems}": "\n".join(
            f"{d['code']} {d['desc']}" for d in data.get("dtc", [])
        )
    }  

    for key, value in replacements.items():  
        replace_all(doc, key, str(value))  

    fill_system_tables(doc, data["faults_raw"])  

    # 🔥 تنسيق جدول معلومات السيارة  
    table = doc.tables[0]  

    for row in table.rows:  
        for i, cell in enumerate(row.cells):  
            # 🔵 العمود الثاني = القيم  
            if i == 1:  
                for p in cell.paragraphs:  
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  
                    for run in p.runs:  
                        run.bold = True  
                        run.font.color.rgb = RGBColor(0, 102, 204)  

            # ⚫ العمود الأول = العناوين  
            else:  
                for p in cell.paragraphs:  
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  
                    for run in p.runs:  
                        run.bold = True  

            # 🔥 توسيط عمودي  
            tc = cell._element  
            tcPr = tc.get_or_add_tcPr()  
            vAlign = OxmlElement('w:vAlign')  
            vAlign.set(qn('w:val'), 'center')  
            tcPr.append(vAlign)  

    doc.save(output_path)